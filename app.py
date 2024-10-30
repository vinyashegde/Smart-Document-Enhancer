from flask import Flask, request, render_template, send_from_directory, redirect, jsonify
from werkzeug.utils import secure_filename
from docx import Document
import PyPDF2
import os
import logging
import google.generativeai as genai

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MODIFIED_FOLDER'] = 'modified'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB upload limit

# Ensure folders exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['MODIFIED_FOLDER'], exist_ok=True)

# Set up logging for debugging
logging.basicConfig(level=logging.DEBUG)

# Set up Google API key
GOOGLE_API_KEY = 'AIzaSyBwXKpzS2syulUlt_OghvgoJkstPsRE70k'  # Replace with your actual API key
genai.configure(api_key=GOOGLE_API_KEY)

# Create a GenerativeModel instance
model = genai.GenerativeModel("gemini-1.5-flash")

# Store debug messages
debug_messages = []

def log_debug(message):
    """Utility function to log debug messages and store them."""
    logging.debug(message)
    debug_messages.append(message)

def extract_text_from_docx(doc_path):
    """Extract text from a .docx file."""
    doc = Document(doc_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_pdf(pdf_path):
    """Extract text from a .pdf file."""
    text = ""
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def generate_content_with_structure(original_content, topic):
    """
    Generate content using Gemini API based on a specified topic while maintaining structure.
    """
    # Create a prompt to maintain structure and formatting
    prompt = (f"From the document which is attached, copy all the content and formatting, "
              f"and generate same-to-same content but with the topic '{topic}'. "
              f"Ensure the structure and length remain the same. "
              f"Content: {original_content}")
    response = model.generate_content([prompt])
    log_debug("Sent prompt to Gemini API.")
    if response:
        log_debug(f"Received modified content of length {len(response.text)}.")
        return response.text
    else:
        log_debug("Failed to get response from Gemini API, returning original content.")
        return original_content

def modify_file_with_length_constraint(file_path, topic):
    log_debug(f"Modifying document: {file_path} with topic: {topic}")

    if not file_path.endswith('.docx'):
        log_debug("Unsupported file type. Only .docx files are supported for this operation.")
        return None

    # Load the document and initialize a new document for modifications
    doc = Document(file_path)
    modified_doc = Document()
    
    in_section = False
    section_content = []
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()

        # Detect the start of "Abstract" or "Introduction" sections
        if text == "Abstract" or text == "1.1 INTRODUCTION":
            # Process and add the previous section's modified content if applicable
            if in_section and section_content:
                modified_content = generate_content_with_structure("\n".join(section_content), topic)
                for paragraph in modified_content.split('\n'):
                    modified_doc.add_paragraph(paragraph)
                section_content = []  # Reset section content for the new section

            # Start a new section
            in_section = True
            current_section = text
            modified_doc.add_paragraph(text)  # Add the section title (e.g., "Abstract" or "1.1 INTRODUCTION")
            continue

        # Detect the end of a section if a new uppercase title starts (assuming a new section)
        elif text.isupper() and in_section:
            # Process the current section content
            modified_content = generate_content_with_structure("\n".join(section_content), topic)
            for paragraph in modified_content.split('\n'):
                modified_doc.add_paragraph(paragraph)
            
            # Reset for the new section
            in_section = False
            section_content = []
            modified_doc.add_paragraph(text)  # Add the new section title as-is
            continue

        # Collect text for the section or copy it directly if not in a section
        if in_section:
            section_content.append(text)
        else:
            modified_doc.add_paragraph(text)

    # Handle the last section if the document ends with it
    if in_section and section_content:
        modified_content = generate_content_with_structure("\n".join(section_content), topic)
        for paragraph in modified_content.split('\n'):
            modified_doc.add_paragraph(paragraph)

    # Save the modified document
    modified_path = os.path.join(app.config['MODIFIED_FOLDER'], 'modified_' + os.path.basename(file_path))
    modified_doc.save(modified_path)
    log_debug(f"Document saved at {modified_path}.")
    return modified_path


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # Clear previous debug messages
        global debug_messages
        debug_messages = []

        # Get the topic from the user
        topic = request.form.get("topic")
        log_debug(f"Received topic: {topic}")
        
        # Check for file upload and topic
        if 'file' not in request.files or not topic:
            log_debug("File or topic missing.")
            return jsonify({"status": "error", "debug": debug_messages})

        file = request.files['file']
        if file.filename == '':
            log_debug("No file selected.")
            return jsonify({"status": "error", "debug": debug_messages})

        if file:
            # Secure and save the uploaded file
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            log_debug(f"File uploaded and saved to {file_path}.")

            # Modify the document based on the specified topic
            modified_file_path = modify_file_with_length_constraint(file_path, topic)

            # Send the modified document for download
            if modified_file_path:
                log_debug(f"Modified file ready for download: {modified_file_path}.")
                return send_from_directory(directory=app.config['MODIFIED_FOLDER'],
                                           path=os.path.basename(modified_file_path),
                                           as_attachment=True)
            else:
                return jsonify({"status": "error", "debug": debug_messages})

    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)
